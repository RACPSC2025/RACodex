"""
Fixtures compartidas para todos los tests de Fénix RAG.

Convenciones:
  - `tmp_path` viene de pytest (directorio temporal por test)
  - Fixtures con `autouse=False` se declaran aquí pero se usan explícitamente
  - Settings de testing se aplican via variables de entorno o model_validate directo
  - No hay I/O real a AWS, Chroma ni PostgreSQL en tests unitarios
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Generator

import pytest


# ─── Configuración de entorno para tests ─────────────────────────────────────

@pytest.fixture(scope="session", autouse=True)
def set_test_environment() -> Generator[None, None, None]:
    """
    Aplica variables de entorno de testing para toda la sesión.

    Evita que los tests lean accidentalmente credenciales reales del .env.
    """
    original = {}
    test_env = {
        "APP_ENV": "testing",
        "AWS_ACCESS_KEY_ID": "test-key-id",
        "AWS_SECRET_ACCESS_KEY": "test-secret",
        "AWS_REGION": "us-east-1",
        "LOG_LEVEL": "DEBUG",
        "DATABASE_URL": "sqlite+aiosqlite:///./test.db",
        "CHROMA_PERSIST_DIR": "/tmp/fenix-test-chroma",
        "BM25_CACHE_DIR": "/tmp/fenix-test-bm25",
        "FLASHRANK_CACHE_DIR": "/tmp/fenix-test-models",
    }

    for key, value in test_env.items():
        original[key] = os.environ.get(key)
        os.environ[key] = value

    yield

    # Restaurar entorno original
    for key, original_value in original.items():
        if original_value is None:
            os.environ.pop(key, None)
        else:
            os.environ[key] = original_value


@pytest.fixture(autouse=True)
def clear_settings_cache() -> Generator[None, None, None]:
    """
    Limpia el caché de settings antes de cada test.

    Garantiza que cada test empiece con settings frescas
    y no haya estado compartido entre tests.
    """
    from src.config.settings import get_settings
    get_settings.cache_clear()
    yield
    get_settings.cache_clear()


@pytest.fixture(autouse=True)
def clear_provider_cache() -> Generator[None, None, None]:
    """Limpia el caché de providers antes de cada test."""
    try:
        from src.config.providers import clear_provider_cache
        clear_provider_cache()
        yield
        clear_provider_cache()
    except ImportError:
        yield


# ─── Fixtures de documentos ───────────────────────────────────────────────────

@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    """Crea un PDF mínimo con texto para tests de ingestion."""
    path = tmp_path / "sample.pdf"
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "ARTÍCULO 1. Objeto del presente decreto.\n" * 20)
        doc.save(str(path))
        doc.close()
    except ImportError:
        # Fallback: escribir bytes mínimos de PDF válido
        path.write_bytes(
            b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj "
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj "
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj "
            b"xref\n0 4\n0000000000 65535 f \n"
            b"trailer<</Root 1 0 R/Size 4>>\n%%EOF"
        )
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Crea un .docx mínimo para tests."""
    path = tmp_path / "sample.docx"
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Contrato de Prestación de Servicios", 0)
        doc.add_paragraph("CLÁUSULA PRIMERA. El contratante se obliga a...")
        doc.save(str(path))
    except ImportError:
        path.write_bytes(b"PK\x03\x04")  # magic bytes ZIP mínimo
    return path


@pytest.fixture
def sample_xlsx_path(tmp_path: Path) -> Path:
    """Crea un .xlsx mínimo para tests."""
    path = tmp_path / "sample.xlsx"
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Normativa"
        ws.append(["Artículo", "Descripción", "Sanción"])
        ws.append(["Art. 1", "Objeto", "N/A"])
        ws.append(["Art. 2", "Ámbito", "Multa"])
        wb.save(str(path))
    except ImportError:
        path.write_bytes(b"PK\x03\x04")
    return path


@pytest.fixture
def legal_text_sample() -> str:
    """Texto legal colombiano de ejemplo para tests de chunking."""
    return """
DECRETO NÚMERO 1072 DE 2015

ARTÍCULO 2.2.4.6.1. Objeto. El presente capítulo tiene por objeto establecer
los parámetros y requisitos mínimos que deben contener los programas de
salud y seguridad en el trabajo.

PARÁGRAFO 1. Las disposiciones del presente decreto se aplican a todos
los empleadores públicos y privados.

ARTÍCULO 2.2.4.6.2. Definiciones. Para los efectos del presente decreto
se aplican las siguientes definiciones:

1. Acción correctiva: Acción tomada para eliminar la causa de una no
conformidad detectada u otra situación indeseable.

2. Acción de mejora: Acción de optimización del Sistema de Gestión de la
Seguridad y Salud en el Trabajo SG-SST, para lograr mejoras en el
desempeño de la organización.

ARTÍCULO 2.2.4.6.3. Política de seguridad y salud en el trabajo. El
empleador o contratante debe establecer por escrito una política de
Seguridad y Salud en el Trabajo – SST.

PARÁGRAFO 2. La política de SST de la empresa debe ser comunicada al
Comité Paritario o Vigía de Seguridad y Salud en el Trabajo.
""".strip()
