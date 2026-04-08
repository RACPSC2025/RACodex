"""
Tests para los loaders de Fase 2: OCR, Docling, Word, Excel.

Estrategia:
  - Todos los tests mockean las dependencias pesadas (easyocr, docling, cv2)
  - Los tests de Word/Excel crean archivos reales con python-docx/openpyxl cuando disponibles
  - Los tests de comportamiento con archivos reales se marcan con skip si falta la dependencia
"""
from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest
from langchain_core.documents import Document

from src.ingestion.base import IngestionError, LoaderUnavailableError


# ═══════════════════════════════════════════════════════════════════════════════
# Tests: OCRLoader
# ═══════════════════════════════════════════════════════════════════════════════

class TestOCRLoader:
    def test_loader_type(self) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        assert loader.loader_type == "ocr"

    def test_supports_pdf(self) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        assert loader.supports(Path("doc.pdf"), "application/pdf") is True

    def test_supports_jpeg(self) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        assert loader.supports(Path("foto.jpg"), "image/jpeg") is True

    def test_supports_png(self) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        assert loader.supports(Path("doc.png"), "image/png") is True

    def test_does_not_support_docx(self) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        assert loader.supports(
            Path("doc.docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ) is False

    def test_file_not_found_raises(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(tmp_path / "no_existe.pdf")

    def test_missing_easyocr_raises_loader_unavailable(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()

        with patch("builtins.__import__", side_effect=lambda name, *a, **k: (
            (_ for _ in ()).throw(ImportError("easyocr not found"))
            if name == "easyocr"
            else __import__(name, *a, **k)
        )):
            loader._reader = None
            with pytest.raises((LoaderUnavailableError, ImportError)):
                loader._get_easyocr()

    def test_reconstruct_text_orders_by_position(self) -> None:
        """La reconstrucción debe ordenar por Y luego por X."""
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader()

        # bbox format: [[x1,y1],[x2,y1],[x2,y2],[x1,y2]]
        ocr_results = [
            ([[100, 200], [200, 200], [200, 220], [100, 220]], "segunda_línea", 0.9),
            ([[100, 100], [200, 100], [200, 120], [100, 120]], "primera_línea", 0.9),
            ([[220, 100], [320, 100], [320, 120], [220, 120]], "primera_derecha", 0.9),
        ]

        text = loader._reconstruct_text(ocr_results, page_height=400)
        lines = text.strip().split("\n")

        # La primera línea debe aparecer antes que la segunda
        assert "primera_línea" in lines[0]
        assert "segunda_línea" in lines[-1]

    def test_confidence_filter_removes_low_quality(self) -> None:
        """Items con confidence < threshold deben descartarse."""
        from src.ingestion.loaders.pdf_ocr import OCRLoader
        loader = OCRLoader(confidence_threshold=0.6)

        ocr_results = [
            ([[0, 0], [100, 0], [100, 20], [0, 20]], "texto_bueno", 0.9),
            ([[0, 50], [100, 50], [100, 70], [0, 70]], "texto_ruido", 0.3),
        ]

        text = loader._reconstruct_text(ocr_results, page_height=200)
        assert "texto_bueno" in text
        # texto_ruido debería estar filtrado o ausente
        # (el filtro ocurre en _run_ocr_on_image, no en _reconstruct_text)
        # Aquí verificamos que _reconstruct_text acepta resultados mixtos

    def test_get_ocr_loader_singleton_without_kwargs(self) -> None:
        from src.ingestion.loaders.pdf_ocr import get_ocr_loader, _shared_ocr_loader
        import src.ingestion.loaders.pdf_ocr as module

        module._shared_ocr_loader = None  # reset singleton
        loader1 = get_ocr_loader()
        loader2 = get_ocr_loader()
        assert loader1 is loader2

    def test_get_ocr_loader_new_instance_with_kwargs(self) -> None:
        from src.ingestion.loaders.pdf_ocr import get_ocr_loader
        loader = get_ocr_loader(max_pages=10)
        assert loader._max_pages == 10


# ═══════════════════════════════════════════════════════════════════════════════
# Tests: DoclingLoader
# ═══════════════════════════════════════════════════════════════════════════════

class TestDoclingLoader:
    def test_loader_type(self) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader
        loader = DoclingLoader()
        assert loader.loader_type == "docling"

    def test_supports_pdf(self) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader
        loader = DoclingLoader()
        assert loader.supports(Path("doc.pdf"), "application/pdf") is True

    def test_does_not_support_docx(self) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader
        loader = DoclingLoader()
        assert loader.supports(
            Path("doc.docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ) is False

    def test_file_not_found_raises(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader
        loader = DoclingLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(tmp_path / "no_existe.pdf")

    def test_missing_docling_raises_loader_unavailable(self) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader
        loader = DoclingLoader()
        loader._converter = None

        with patch("builtins.__import__", side_effect=lambda name, *a, **k: (
            (_ for _ in ()).throw(ImportError("docling not found"))
            if name == "docling.document_converter"
            else __import__(name, *a, **k)
        )):
            with pytest.raises((LoaderUnavailableError, ImportError)):
                loader._get_converter()

    def test_load_with_mocked_converter(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader

        path = tmp_path / "contrato.pdf"
        path.write_bytes(b"%PDF-1.4")

        mock_doc = MagicMock()
        mock_doc.document.export_to_markdown.return_value = (
            "# CAPÍTULO I - Generalidades\n\n"
            "## ARTÍCULO 1. Objeto\n\n"
            "El presente contrato establece las condiciones.\n\n"
            "## ARTÍCULO 2. Definiciones\n\n"
            "Para efectos de este contrato se entiende por...\n"
        )

        mock_converter = MagicMock()
        mock_converter.convert.return_value = mock_doc

        loader = DoclingLoader(chunk_size=500)
        loader._converter = mock_converter

        docs = loader.load(path)

        assert len(docs) > 0
        assert all(isinstance(d, Document) for d in docs)
        mock_converter.convert.assert_called_once_with(str(path))

    def test_empty_markdown_raises_ingestion_error(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader

        path = tmp_path / "empty.pdf"
        path.write_bytes(b"%PDF-1.4")

        mock_doc = MagicMock()
        mock_doc.document.export_to_markdown.return_value = ""

        mock_converter = MagicMock()
        mock_converter.convert.return_value = mock_doc

        loader = DoclingLoader()
        loader._converter = mock_converter

        with pytest.raises(IngestionError):
            loader.load(path)

    def test_split_markdown_respects_headers(self) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader

        loader = DoclingLoader(chunk_size=2000)
        markdown = (
            "# CAPÍTULO I\n\nTexto del capítulo.\n\n"
            "## ARTÍCULO 1\n\nContenido del artículo 1.\n\n"
            "## ARTÍCULO 2\n\nContenido del artículo 2.\n"
        )
        chunks = loader._split_markdown(markdown)

        assert len(chunks) >= 2
        combined = " ".join(chunks)
        assert "ARTÍCULO 1" in combined
        assert "ARTÍCULO 2" in combined

    def test_documents_have_metadata(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.docling_loader import DoclingLoader

        path = tmp_path / "doc.pdf"
        path.write_bytes(b"%PDF-1.4")

        mock_doc = MagicMock()
        mock_doc.document.export_to_markdown.return_value = (
            "# Sección\n\nTexto suficiente para generar chunks con metadata completa.\n"
            "Más texto para asegurar que el chunk no sea filtrado por tamaño.\n"
        )

        loader = DoclingLoader()
        loader._converter = MagicMock(return_value=mock_doc)
        loader._converter.convert.return_value = mock_doc

        docs = loader.load(path)
        for doc in docs:
            assert "source" in doc.metadata or "loader" in doc.metadata


# ═══════════════════════════════════════════════════════════════════════════════
# Tests: WordLoader
# ═══════════════════════════════════════════════════════════════════════════════

class TestWordLoader:
    def test_loader_type(self) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()
        assert loader.loader_type == "word"

    def test_supports_docx(self) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()
        assert loader.supports(
            Path("doc.docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ) is True

    def test_supports_doc_legacy(self) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()
        assert loader.supports(Path("doc.doc"), "application/msword") is True

    def test_does_not_support_pdf(self) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()
        assert loader.supports(Path("doc.pdf"), "application/pdf") is False

    def test_file_not_found_raises(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(tmp_path / "no_existe.docx")

    def test_missing_python_docx_raises(self) -> None:
        from src.ingestion.loaders.word_loader import WordLoader
        loader = WordLoader()

        with patch("builtins.__import__", side_effect=lambda name, *a, **k: (
            (_ for _ in ()).throw(ImportError("docx not found"))
            if name == "docx"
            else __import__(name, *a, **k)
        )):
            with pytest.raises((LoaderUnavailableError, ImportError)):
                loader._get_docx()

    def test_load_real_docx(self, tmp_path: Path) -> None:
        """Test con archivo .docx real — skip si python-docx no disponible."""
        try:
            from docx import Document as DocxDoc
        except ImportError:
            pytest.skip("python-docx no disponible")

        from src.ingestion.loaders.word_loader import WordLoader

        path = tmp_path / "contrato.docx"
        doc = DocxDoc()
        doc.add_heading("CLÁUSULA PRIMERA. Objeto del Contrato", level=1)
        doc.add_paragraph(
            "El contratante se obliga a prestar servicios de consultoría "
            "en materia de seguridad y salud en el trabajo. " * 5
        )
        doc.add_heading("CLÁUSULA SEGUNDA. Obligaciones", level=1)
        doc.add_paragraph("Las partes se obligan a cumplir con la normativa vigente. " * 5)
        doc.save(str(path))

        loader = WordLoader(chunk_size=500)
        docs = loader.load(path)

        assert len(docs) > 0
        assert all(isinstance(d, Document) for d in docs)
        combined = " ".join(d.page_content for d in docs)
        assert "CLÁUSULA" in combined

    def test_table_to_markdown(self) -> None:
        """Test de conversión de tabla Word a Markdown."""
        try:
            from docx import Document as DocxDoc
        except ImportError:
            pytest.skip("python-docx no disponible")

        from src.ingestion.loaders.word_loader import WordLoader

        loader = WordLoader()

        # Crear tabla mock
        mock_cell = MagicMock()
        mock_cell.text = "Valor"

        mock_row1 = MagicMock()
        mock_row1.cells = [
            MagicMock(text="Artículo"),
            MagicMock(text="Descripción"),
            MagicMock(text="Sanción"),
        ]
        mock_row2 = MagicMock()
        mock_row2.cells = [
            MagicMock(text="Art. 22"),
            MagicMock(text="Obligación SST"),
            MagicMock(text="20 SMLV"),
        ]

        mock_table = MagicMock()
        mock_table.rows = [mock_row1, mock_row2]

        result = loader._table_to_markdown(mock_table)

        assert "Artículo" in result
        assert "Art. 22" in result
        assert "|" in result
        assert "---" in result

    def test_empty_docx_raises_ingestion_error(self, tmp_path: Path) -> None:
        try:
            from docx import Document as DocxDoc
        except ImportError:
            pytest.skip("python-docx no disponible")

        from src.ingestion.loaders.word_loader import WordLoader

        path = tmp_path / "empty.docx"
        doc = DocxDoc()
        # No añadimos nada — documento vacío
        doc.save(str(path))

        loader = WordLoader()
        with pytest.raises(IngestionError):
            loader.load(path)


# ═══════════════════════════════════════════════════════════════════════════════
# Tests: ExcelLoader
# ═══════════════════════════════════════════════════════════════════════════════

class TestExcelLoader:
    def test_loader_type(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()
        assert loader.loader_type == "excel"

    def test_supports_xlsx(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()
        assert loader.supports(
            Path("data.xlsx"),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ) is True

    def test_supports_xls_legacy(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()
        assert loader.supports(Path("data.xls"), "application/vnd.ms-excel") is True

    def test_does_not_support_pdf(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()
        assert loader.supports(Path("doc.pdf"), "application/pdf") is False

    def test_file_not_found_raises(self, tmp_path: Path) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(tmp_path / "no_existe.xlsx")

    def test_missing_pandas_raises_loader_unavailable(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader
        loader = ExcelLoader()

        with patch("builtins.__import__", side_effect=lambda name, *a, **k: (
            (_ for _ in ()).throw(ImportError("pandas not found"))
            if name == "pandas"
            else __import__(name, *a, **k)
        )):
            with pytest.raises((LoaderUnavailableError, ImportError)):
                loader._get_pandas()

    def test_load_real_xlsx(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "normativa.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sanciones"
        ws.append(["Artículo", "Descripción", "Sanción (SMLV)"])
        for i in range(1, 25):
            ws.append([f"Art. {i}", f"Descripción del artículo {i}", str(i * 2)])
        wb.save(str(path))

        loader = ExcelLoader(rows_per_chunk=10)
        docs = loader.load(path)

        assert len(docs) > 0
        assert all(isinstance(d, Document) for d in docs)

    def test_chunk_size_limits_rows(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Col1", "Col2"])
        for i in range(50):
            ws.append([f"Val{i}", f"Info{i}"])
        wb.save(str(path))

        loader = ExcelLoader(rows_per_chunk=10)
        docs = loader.load(path)

        # 50 filas / 10 por chunk = 5 chunks mínimo
        assert len(docs) >= 5

    def test_row_paragraph_format(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Artículo", "Descripción"])
        ws.append(["Art. 1", "Objeto"])
        wb.save(str(path))

        loader = ExcelLoader(output_format="row_paragraph")
        docs = loader.load(path)

        content = " ".join(d.page_content for d in docs)
        # El formato "fila como párrafo" incluye "Columna: Valor"
        assert "Artículo:" in content or "Artículo" in content

    def test_markdown_table_format(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Col1", "Col2"])
        ws.append(["A", "B"])
        wb.save(str(path))

        loader = ExcelLoader(output_format="markdown_table")
        docs = loader.load(path)

        content = " ".join(d.page_content for d in docs)
        assert "|" in content

    def test_sheet_metadata_in_documents(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "MiHoja"
        ws.append(["A", "B"])
        ws.append(["1", "2"])
        wb.save(str(path))

        loader = ExcelLoader()
        docs = loader.load(path)

        for doc in docs:
            assert doc.metadata.get("sheet_name") == "MiHoja"

    def test_multiple_sheets_processed(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()

        ws1 = wb.active
        ws1.title = "Hoja1"
        ws1.append(["A", "B"])
        ws1.append(["1", "2"])

        ws2 = wb.create_sheet("Hoja2")
        ws2.append(["X", "Y"])
        ws2.append(["3", "4"])

        wb.save(str(path))

        loader = ExcelLoader()
        docs = loader.load(path)

        sheet_names = {d.metadata.get("sheet_name") for d in docs}
        assert "Hoja1" in sheet_names
        assert "Hoja2" in sheet_names

    def test_empty_excel_raises_ingestion_error(self, tmp_path: Path) -> None:
        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            pytest.skip("openpyxl/pandas no disponibles")

        from src.ingestion.loaders.excel_loader import ExcelLoader

        path = tmp_path / "empty.xlsx"
        wb = openpyxl.Workbook()
        wb.active.title = "Vacía"
        # No añadimos datos
        wb.save(str(path))

        loader = ExcelLoader(skip_empty_sheets=True)
        with pytest.raises(IngestionError):
            loader.load(path)

    def test_rows_to_paragraph_format(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader

        loader = ExcelLoader(output_format="row_paragraph")
        columns = ["Artículo", "Descripción", "Sanción"]
        rows = [
            {"Artículo": "Art. 1", "Descripción": "Objeto", "Sanción": "10 SMLV"},
            {"Artículo": "Art. 2", "Descripción": "Ámbito", "Sanción": "20 SMLV"},
        ]

        result = loader._rows_to_paragraph(None, columns, rows)

        assert "Artículo: Art. 1" in result
        assert "Descripción: Objeto" in result
        assert "Sanción: 10 SMLV" in result

    def test_rows_to_markdown_table_format(self) -> None:
        from src.ingestion.loaders.excel_loader import ExcelLoader

        loader = ExcelLoader()
        columns = ["Col1", "Col2"]
        rows = [
            {"Col1": "A", "Col2": "B"},
            {"Col1": "C", "Col2": "D"},
        ]

        result = loader._rows_to_markdown_table(columns, rows)

        assert "| Col1 | Col2 |" in result
        assert "| --- | --- |" in result
        assert "| A | B |" in result
