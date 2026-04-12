"""
Loader de documentos Word (.docx / .doc legacy).

Estrategia de extracción para documentos legales en Word:
  1. Párrafos con estilos de Heading → equivalentes a ARTÍCULO/CAPÍTULO
  2. Párrafos normales → texto del artículo
  3. Tablas → serialización en formato | col | col | (Markdown GFM)
  4. Listas numeradas/con viñetas → texto con prefijo numérico/guión

Por qué preservar estilos:
  Un contrato en Word usa "Heading 1" para CLÁUSULA PRIMERA y
  "Normal" para el cuerpo. Si ignoramos los estilos y tratamos
  todo como texto plano, el chunker no puede encontrar límites naturales.

Soporte .doc legacy:
  python-docx NO soporta .doc (formato binario de Word 97-2003).
  Para .doc se intenta conversión via LibreOffice si está disponible.
  Si no hay LibreOffice, se lanza LoaderUnavailableError con instrucciones.

Requiere: python-docx >= 1.1
  pip install python-docx
"""

from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from langchain_core.documents import Document

from src.config.logging import get_logger
from src.config.settings import get_settings
from src.ingestion.base import BaseLoader, IngestionError, LoaderUnavailableError
from src.ingestion.processors.hierarchical_chunker import HierarchicalChunker, get_hierarchical_chunker
from src.ingestion.processors.text_cleaner import get_cleaner

log = get_logger(__name__)

# Estilos de Word que actúan como separadores de sección en documentos legales
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3",
    "título 1", "título 2", "título 3",    # español
    "estilo1", "estilo2",
}

# Prefijo de lista para numerales legales
_LIST_PREFIX = {
    "list paragraph", "list bullet", "list number",
    "párrafo de lista", "viñeta de lista",
}


class WordLoader(BaseLoader):
    """
    Loader de documentos Word (.docx) para normativa y contratos legales.

    Extrae texto preservando la estructura jerárquica del documento
    (headings, párrafos, tablas) antes de pasar al chunker.
    """

    def __init__(
        self,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
        cleaner_profile: str = "default",
        include_tables: bool = True,
        include_headers_footers: bool = False,
    ) -> None:
        settings = get_settings()
        self._chunk_size = chunk_size or settings.chunk_size
        self._chunk_overlap = chunk_overlap or settings.chunk_overlap
        self._cleaner_profile = cleaner_profile
        self._include_tables = include_tables
        self._include_headers_footers = include_headers_footers
        self._chunker: LegalChunker | None = None

    @property
    def loader_type(self) -> str:
        return "word"

    def supports(self, path: Path, mime_type: str) -> bool:
        return mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }

    def _get_docx(self) -> object:
        try:
            import docx  # noqa: PLC0415
            return docx
        except ImportError as exc:
            raise LoaderUnavailableError(
                "python-docx no está instalado.\n"
                "Ejecuta: pip install python-docx"
            ) from exc

    def _get_chunker(self) -> LegalChunker:
        if self._chunker is None:
            self._chunker = LegalChunker(
                chunk_size=self._chunk_size,
                chunk_overlap=self._chunk_overlap,
            )
        return self._chunker

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        """
        Convierte .doc legacy a .docx usando LibreOffice en modo headless.

        Requiere LibreOffice instalado:
          Ubuntu/Debian: sudo apt-get install libreoffice
          macOS: brew install libreoffice

        Returns:
            Path al .docx generado en directorio temporal.

        Raises:
            LoaderUnavailableError: Si LibreOffice no está instalado.
            IngestionError: Si la conversión falla.
        """
        try:
            result = subprocess.run(
                ["libreoffice", "--version"],
                capture_output=True,
                timeout=10,
            )
            if result.returncode != 0:
                raise FileNotFoundError("LibreOffice no disponible")
        except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
            raise LoaderUnavailableError(
                "El formato .doc (Word 97-2003) requiere LibreOffice para conversión.\n"
                "Ubuntu/Debian: sudo apt-get install libreoffice\n"
                "macOS: brew install libreoffice\n"
                "Alternativa: convierte el archivo a .docx manualmente."
            ) from exc

        tmp_dir = Path(tempfile.mkdtemp())
        log.info("converting_doc_to_docx", file=doc_path.name, tmp_dir=str(tmp_dir))

        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", str(tmp_dir),
                    str(doc_path),
                ],
                capture_output=True,
                timeout=60,
            )
        except subprocess.TimeoutExpired as exc:
            raise IngestionError(
                f"Timeout convirtiendo {doc_path.name} con LibreOffice",
                path=doc_path,
                cause=exc,
            ) from exc

        if result.returncode != 0:
            raise IngestionError(
                f"LibreOffice falló convirtiendo {doc_path.name}: "
                f"{result.stderr.decode()[:200]}",
                path=doc_path,
            )

        docx_path = tmp_dir / f"{doc_path.stem}.docx"
        if not docx_path.exists():
            raise IngestionError(
                f"LibreOffice no generó el .docx esperado en {tmp_dir}",
                path=doc_path,
            )

        return docx_path

    def _extract_text_from_docx(self, path: Path) -> str:
        """
        Extrae texto estructurado de un .docx.

        Convierte la jerarquía de estilos Word en marcadores de texto
        que el LegalChunker puede reconocer como límites de sección.
        """
        docx = self._get_docx()

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise IngestionError(
                f"No se pudo abrir el documento Word: {path.name}",
                path=path,
                cause=exc,
            ) from exc

        parts: list[str] = []

        # ── Párrafos del cuerpo principal ─────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()

            if style_name in _HEADING_STYLES:
                # Heading → forzar salto de línea para que el chunker lo detecte
                parts.append(f"\n\n{text}\n")
            elif style_name in _LIST_PREFIX:
                parts.append(f"  - {text}")
            else:
                parts.append(text)

        # ── Tablas ────────────────────────────────────────────────────────────
        if self._include_tables:
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    parts.append(f"\n{table_md}\n")

        return "\n".join(parts)

    def _table_to_markdown(self, table: object) -> str:
        """
        Convierte una tabla Word a Markdown GFM.

        Formato:
          | Col1 | Col2 | Col3 |
          |------|------|------|
          | Val1 | Val2 | Val3 |
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        if len(rows) == 1:
            # Solo encabezado — retornar como lista
            return " | ".join(rows[0])

        # Encabezado + separador + datos
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
        data_rows = [
            "| " + " | ".join(row) + " |"
            for row in rows[1:]
        ]

        return "\n".join([header, separator, *data_rows])

    def load(self, path: Path) -> list[Document]:
        """
        Carga un documento Word y retorna sus chunks como Documents.

        Args:
            path: Ruta al .docx o .doc. Debe existir.

        Returns:
            Lista de Documents con texto estructurado y metadata.

        Raises:
            FileNotFoundError: Si el archivo no existe.
            IngestionError: Si el documento no puede procesarse.
            LoaderUnavailableError: Si python-docx no está instalado,
                                   o si .doc requiere LibreOffice.
        """
        path = Path(path).resolve()

        if not path.exists():
            raise FileNotFoundError(f"Documento no encontrado: {path}")

        log.info("word_loading", file=path.name)

        # Convertir .doc legacy si es necesario
        effective_path = path
        is_legacy = path.suffix.lower() == ".doc"
        if is_legacy:
            effective_path = self._convert_doc_to_docx(path)
            log.info("word_doc_converted", original=path.name, docx=effective_path.name)

        # Extraer texto con estructura preservada
        raw_text = self._extract_text_from_docx(effective_path)

        # Limpiar directorio temporal si hubo conversión
        if is_legacy:
            try:
                effective_path.unlink(missing_ok=True)
                effective_path.parent.rmdir()
            except Exception:
                pass  # no crítico

        if not raw_text.strip():
            raise IngestionError(
                f"El documento '{path.name}' no contiene texto extraíble.",
                path=path,
            )

        # Limpiar y chunkear
        chunker = self._get_chunker()
        documents = chunker.chunk_with_profile(
            text=raw_text,
            source_path=path,
            loader_type=self.loader_type,
            cleaner_profile=self._cleaner_profile,
            add_header=True,
        )

        log.info(
            "word_loaded",
            file=path.name,
            chars=len(raw_text),
            chunks=len(documents),
        )

        return documents


def get_word_loader(**kwargs) -> WordLoader:
    """Factory function para WordLoader."""
    return WordLoader(**kwargs)
